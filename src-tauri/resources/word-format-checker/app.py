from __future__ import annotations

import re
import hashlib
from io import BytesIO

import pandas as pd
import streamlit as st
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

from checker import check_docx, DEFAULT_RULES, PAPER_SIZES_CM, CN_FONT_SIZES_PT


st.set_page_config(page_title="Word 格式检查工具", layout="wide")
st.title("📄 Word 格式检查工具 (v0.9)")
st.warning("⚠️ 本工具结果仅供辅助复核，不能替代人工审校；正式提交前请务必人工逐项复查。")

with st.expander("查看使用说明", expanded=False):
    try:
        with open("README.md", "r", encoding="utf-8") as f:
            st.markdown(f.read())
    except Exception as e:
        st.warning(f"使用说明加载失败：{e}")

paper_options = [f"{k} ({v[0]:.2f}×{v[1]:.2f} cm)" for k, v in PAPER_SIZES_CM.items()]
paper_keys = list(PAPER_SIZES_CM.keys())

size_labels = [f"{k} ({v:g} pt)" for k, v in CN_FONT_SIZES_PT.items()]
size_map = {f"{k} ({v:g} pt)": k for k, v in CN_FONT_SIZES_PT.items()}

common_fonts = ["宋体", "仿宋", "黑体", "楷体", "微软雅黑", "等线", "Times New Roman", "Arial", "Calibri", "自定义..."]


def aggregate_rows(rows: list[dict]):
    """展示层合并：同一位置+类别+规则合并为一条，不影响原始检测结果。"""
    grouped = {}
    for i, r in enumerate(rows, start=1):
        key = (str(r.get("location", "")), str(r.get("category", "")), str(r.get("rule", "")))
        if key not in grouped:
            grouped[key] = {
                "row": dict(r),
                "count": 1,
                "issue_ids": [i],
            }
        else:
            grouped[key]["count"] += 1
            grouped[key]["issue_ids"].append(i)

    merged = []
    for _, v in grouped.items():
        row = v["row"]
        cnt = v["count"]
        ids = v["issue_ids"]
        if cnt > 1:
            row["actual"] = f"{row.get('actual','')}（同类命中 {cnt} 次）"
        row["_issue_ids"] = ids
        row["_count"] = cnt
        merged.append(row)
    return merged


def build_annotated_docx(file_bytes: bytes, rows: list[dict]):
    doc = Document(BytesIO(file_bytes))

    def mark_paragraph(p, msg: str, force_visible: bool = False):
        # 对空段落强制写入可见标记，避免“看起来没标上”
        if force_visible and not (p.text or "").strip():
            r = p.add_run(f"【{msg}：原为空段落】")
        else:
            r = p.add_run(f"  【{msg}】")
        r.bold = True
        r.font.highlight_color = WD_COLOR_INDEX.YELLOW

    max_marks = 2000
    mark_count = 0
    unmarked_ids = []

    # 页首汇总段（用于承接结构级/无法精确定位的问题）
    top_para = doc.paragraphs[0] if doc.paragraphs else doc.add_paragraph("")
    top_header = top_para.add_run("【自动标记汇总】以下问题为结构级或无法精确定位，按问题编号回查：")
    top_header.bold = True
    top_header.font.highlight_color = WD_COLOR_INDEX.YELLOW

    for i, row in enumerate(rows, start=1):
        if mark_count >= max_marks:
            unmarked_ids.append(i)
            continue
        loc = str(row.get("location", ""))
        rule = str(row.get("rule", ""))
        level = str(row.get("level", ""))
        msg = f"问题#{i} {level}/{rule}"
        force_visible = (rule == "空段落")

        m_para = re.search(r"段落#(\d+)", loc)
        if m_para:
            idx = int(m_para.group(1))
            if 1 <= idx <= len(doc.paragraphs):
                mark_paragraph(doc.paragraphs[idx - 1], msg, force_visible=force_visible)
                mark_count += 1
                continue

        m_cell = re.search(r"表格#(\d+).*行(\d+)列(\d+)", loc)
        if m_cell:
            t_idx, r_idx, c_idx = map(int, m_cell.groups())
            if 1 <= t_idx <= len(doc.tables):
                tbl = doc.tables[t_idx - 1]
                if 1 <= r_idx <= len(tbl.rows) and 1 <= c_idx <= len(tbl.rows[r_idx - 1].cells):
                    cell = tbl.rows[r_idx - 1].cells[c_idx - 1]
                    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph("")
                    mark_paragraph(p, msg)
                    mark_count += 1
                    continue

        # 回退策略：按文本片段在正文中搜索首个命中并标记
        snippet = str(row.get("snippet", "")).strip()
        if snippet:
            needle = snippet[:20]
            matched = False
            for p in doc.paragraphs:
                if needle and needle in (p.text or ""):
                    mark_paragraph(p, msg)
                    mark_count += 1
                    matched = True
                    break
            if matched:
                continue

        # 仍无法定位：插入页首汇总，至少保证“可见”
        top_note = top_para.add_run(f"\n- {msg}")
        top_note.font.highlight_color = WD_COLOR_INDEX.YELLOW
        mark_count += 1
        unmarked_ids.append(i)

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue(), mark_count, unmarked_ids

with st.sidebar:
    st.header("规则配置")

    with st.expander("页面", expanded=True):
        default_paper_idx = paper_keys.index(DEFAULT_RULES["paper"]) if DEFAULT_RULES["paper"] in paper_keys else 0
        paper_label = st.selectbox("纸张（含尺寸）", options=paper_options, index=default_paper_idx)
        paper = paper_label.split(" ")[0]

        c1, c2 = st.columns(2)
        with c1:
            margin_top_cm = st.number_input("上边距(cm)", 0.0, 10.0, float(DEFAULT_RULES["margin_top_cm"]))
            margin_left_cm = st.number_input("左边距(cm)", 0.0, 10.0, float(DEFAULT_RULES["margin_left_cm"]))
        with c2:
            margin_bottom_cm = st.number_input("下边距(cm)", 0.0, 10.0, float(DEFAULT_RULES["margin_bottom_cm"]))
            margin_right_cm = st.number_input("右边距(cm)", 0.0, 10.0, float(DEFAULT_RULES["margin_right_cm"]))

        max_pages = st.number_input("页数上限", min_value=1, max_value=1000, value=int(DEFAULT_RULES["max_pages"]))

    with st.expander("正文（全篇统一）", expanded=True):
        default_font_idx = common_fonts.index(DEFAULT_RULES["font_name"]) if DEFAULT_RULES["font_name"] in common_fonts else 0
        font_choice = st.selectbox("字体（常用）", common_fonts, index=default_font_idx)
        font_name = st.text_input("自定义字体", DEFAULT_RULES["font_name"]) if font_choice == "自定义..." else font_choice
        font_name_alt = st.text_input("字体别名（可选）", DEFAULT_RULES["font_name_alt"])

        default_size_label = f"四号 ({CN_FONT_SIZES_PT['四号']:g} pt)"
        size_label = st.selectbox("字号（中文习惯）", options=size_labels, index=size_labels.index(default_size_label) if default_size_label in size_labels else 0)
        font_size_cn = size_map[size_label]

        c3, c4 = st.columns(2)
        with c3:
            line_spacing = st.number_input("固定行距(pt)", 10.0, 80.0, float(DEFAULT_RULES["line_spacing_pt"]))
            first_indent = st.number_input("首行缩进(字符)", 0.0, 6.0, float(DEFAULT_RULES["first_line_indent_chars"]))
        with c4:
            space_before = st.number_input("段前距(pt)", 0.0, 50.0, float(DEFAULT_RULES["space_before_pt"]))
            space_after = st.number_input("段后距(pt)", 0.0, 50.0, float(DEFAULT_RULES["space_after_pt"]))

    with st.expander("允许项", expanded=False):
        c5, c6 = st.columns(2)
        with c5:
            allow_bold = st.checkbox("允许加粗", DEFAULT_RULES["allow_bold"])
            allow_non_black_text = st.checkbox("允许非黑色文字", DEFAULT_RULES["allow_non_black_text"])
            allow_empty_para = st.checkbox("允许空段落", DEFAULT_RULES["allow_empty_para"])
            allow_header = st.checkbox("允许页眉", DEFAULT_RULES["allow_header"])
            allow_toc = st.checkbox("允许目录", DEFAULT_RULES["allow_toc"])
        with c6:
            allow_italic = st.checkbox("允许斜体", DEFAULT_RULES["allow_italic"])
            allow_underline = st.checkbox("允许下划线", DEFAULT_RULES["allow_underline"])
            allow_leading_space = st.checkbox("允许段首空格缩进", DEFAULT_RULES["allow_leading_space"])
            allow_any_space = st.checkbox("允许出现任意空格", DEFAULT_RULES["allow_any_space"])
            allow_page_number = st.checkbox("允许页码", DEFAULT_RULES["allow_page_number"])
            allow_cover = st.checkbox("允许封面", DEFAULT_RULES["allow_cover"])

    with st.expander("表格 / 对象", expanded=False):
        c7, c8 = st.columns(2)
        with c7:
            allow_table_row_break = st.checkbox("允许表格跨页断行", DEFAULT_RULES["allow_table_row_break"])
            require_table_border = st.checkbox("要求表格有边框", DEFAULT_RULES["require_table_border"])
            align_labels = ["左对齐", "居中", "右对齐"]
            align_map = {"左对齐": "left", "居中": "center", "右对齐": "right"}
            reverse_align_map = {v: k for k, v in align_map.items()}
            table_text_align_label = st.selectbox("表内文字水平对齐", align_labels, index=align_labels.index(reverse_align_map.get(DEFAULT_RULES["table_text_align"], "左对齐")))
            table_text_align = align_map[table_text_align_label]
        with c8:
            allow_floating_shapes = st.checkbox("允许浮动对象", DEFAULT_RULES["allow_floating_shapes"])
            valign_labels = ["顶端", "居中", "底端"]
            valign_map = {"顶端": "top", "居中": "center", "底端": "bottom"}
            reverse_valign_map = {v: k for k, v in valign_map.items()}
            table_vertical_align_label = st.selectbox("单元格垂直对齐", valign_labels, index=valign_labels.index(reverse_valign_map.get(DEFAULT_RULES["table_vertical_align"], "居中")))
            table_vertical_align = valign_map[table_vertical_align_label]
            table_first_line_indent_chars = st.number_input("表内首行缩进(字符)", min_value=0.0, max_value=6.0, value=float(DEFAULT_RULES["table_first_line_indent_chars"]), step=0.5)

    forbidden_chars = st.text_input("禁用字符集合", DEFAULT_RULES["forbidden_chars"])
    supplier_keywords = st.text_area("供应商身份关键词（逗号分隔）", DEFAULT_RULES["supplier_keywords"], height=80)

    with st.expander("身份信息检测", expanded=False):
        pii_level_label = st.selectbox("检测强度", ["高（默认）", "中", "低"], index=0)
        pii_level_map = {"高（默认）": "high", "中": "medium", "低": "low"}
        pii_level = pii_level_map[pii_level_label]
        pii_high_sensitive_default = st.checkbox("高敏感项默认按错误输出（手机号/邮箱/证件号）", DEFAULT_RULES["pii_high_sensitive_default"])

    rules = {
        "paper": paper,
        "margin_top_cm": margin_top_cm,
        "margin_bottom_cm": margin_bottom_cm,
        "margin_left_cm": margin_left_cm,
        "margin_right_cm": margin_right_cm,
        "max_pages": int(max_pages),
        "font_name": font_name,
        "font_name_alt": font_name_alt,
        "font_size_pt": font_size_cn,
        "line_spacing_pt": line_spacing,
        "first_line_indent_chars": first_indent,
        "space_before_pt": space_before,
        "space_after_pt": space_after,
        "allow_bold": allow_bold,
        "allow_italic": allow_italic,
        "allow_non_black_text": allow_non_black_text,
        "allow_underline": allow_underline,
        "allow_empty_para": allow_empty_para,
        "allow_leading_space": allow_leading_space,
        "allow_any_space": allow_any_space,
        "allow_header": allow_header,
        "allow_page_number": allow_page_number,
        "allow_toc": allow_toc,
        "allow_cover": allow_cover,
        "allow_table_row_break": allow_table_row_break,
        "require_table_border": require_table_border,
        "table_text_align": table_text_align,
        "table_vertical_align": table_vertical_align,
        "table_first_line_indent_chars": table_first_line_indent_chars,
        "allow_floating_shapes": allow_floating_shapes,
        "forbidden_chars": forbidden_chars,
        "supplier_keywords": supplier_keywords,
        "pii_level": pii_level,
        "pii_high_sensitive_default": pii_high_sensitive_default,
    }

uploaded = st.file_uploader("上传 Word 文档（.docx）", type=["docx"], help="仅支持 .docx，暂不支持 .doc")
if uploaded is None:
    st.info("请先上传一个 Word 文档（.docx）。")
    st.stop()

uploaded_bytes = uploaded.getvalue()
uploaded_digest = hashlib.sha256(uploaded_bytes).hexdigest()

if "rows" not in st.session_state:
    st.session_state.rows = None
if "annotated" not in st.session_state:
    st.session_state.annotated = None
if "last_upload_digest" not in st.session_state:
    st.session_state.last_upload_digest = None

# 避免“更换文档后沿用旧检测结果”导致展示和标记错位
if st.session_state.last_upload_digest != uploaded_digest:
    st.session_state.rows = None
    st.session_state.annotated = None
    st.session_state.last_upload_digest = uploaded_digest
b1, b2 = st.columns(2)
with b1:
    do_check = st.button("开始检查", use_container_width=True, type="primary")
with b2:
    do_mark = st.button("生成标记副本", use_container_width=True)

if do_check:
    with st.spinner("正在检查文档格式，请稍候..."):
        try:
            violations = check_docx(uploaded_bytes, rules)
            st.session_state.rows = [v.to_dict() for v in violations]
            st.session_state.annotated = None
        except Exception as e:
            st.error(f"文档解析失败：{e}")

rows = st.session_state.rows
if rows is None:
    st.warning("请先点击“开始检查”。")
    st.stop()

merged_rows = aggregate_rows(rows)

# 先处理“生成标记副本”，让下载按钮始终紧跟动作按钮显示
if do_mark:
    try:
        # 标记副本仍基于“原始检测结果”逐条编号，确保不丢失检测效果
        annotated, marked_count, unmarked_ids = build_annotated_docx(uploaded_bytes, rows)
        st.session_state.annotated = annotated
        st.info(f"已标记 {marked_count} 条。未能自动定位 {len(unmarked_ids)} 条（多为页面级/结构级规则）。")
        if unmarked_ids:
            preview_ids = ", ".join([f"问题#{x}" for x in unmarked_ids[:30]])
            st.caption(f"未标记问题编号（前30条）：{preview_ids}")
    except Exception as e:
        st.error(f"生成标记副本失败：{e}")

if st.session_state.annotated is not None:
    st.download_button(
        "下载标记副本(.docx)",
        st.session_state.annotated,
        "format_check_标记副本.docx",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        use_container_width=True,
    )

st.success(f"检查完成：原始命中 {len(rows)} 条；合并展示 {len(merged_rows)} 条")
st.caption("说明：表格为“合并展示”（同位置+同规则归并）；标记副本仍按原始命中逐条编号。可用‘问题编号范围/原始编号明细’与【问题#N】互相映射。")

if len(merged_rows) == 0:
    st.balloons()
    st.write("🎉 未发现格式问题（按当前规则）。")
else:
    df = pd.DataFrame(merged_rows)

    format_categories = {"页面", "页眉", "页脚", "段落", "字体", "字符样式", "文本", "结构", "表格", "图片"}

    def _biz_level(row):
        cat = str(row.get("category", ""))
        rule = str(row.get("rule", ""))
        if cat in format_categories:
            return "必改"
        if cat == "暗标合规" or "身份信息" in rule:
            return "警告"
        return "必改"

    def _raw_level_text(v):
        return {"error": "error", "warn": "warn", "info": "info"}.get(str(v), str(v))

    df["raw_level"] = df["level"].apply(_raw_level_text) if "level" in df.columns else "warn"
    df["biz_level"] = df.apply(_biz_level, axis=1)

    c1, c2, c3 = st.columns(3)
    c1.metric("问题总数", int(len(df)))
    c2.metric("必改项", int((df["biz_level"] == "必改").sum()))
    c3.metric("警告项", int((df["biz_level"] == "警告").sum()))

    def _issue_label(ids):
        if not ids:
            return "-"
        if len(ids) == 1:
            return f"问题#{ids[0]}"
        return f"问题#{ids[0]}~#{ids[-1]}"

    def _issue_ids_detail(ids):
        if not ids:
            return "-"
        if len(ids) <= 8:
            return ", ".join([f"#{x}" for x in ids])
        return ", ".join([f"#{x}" for x in ids[:8]]) + f" ... 共{len(ids)}条"

    ids_col = df.get("_issue_ids", pd.Series([[] for _ in range(len(df))]))
    df.insert(0, "issue_id", ids_col.apply(_issue_label))
    df.insert(1, "issue_ids_detail", ids_col.apply(_issue_ids_detail))

    df = df.rename(columns={
        "issue_id": "问题编号范围",
        "issue_ids_detail": "原始编号明细",
        "_count": "命中次数",
        "raw_level": "原始级别",
        "biz_level": "处理级别",
        "category": "类别",
        "rule": "规则",
        "location": "高级定位信息",
        "expected": "期望值",
        "actual": "实际值",
        "snippet": "文本片段",
        "suggestion": "修复建议",
    })

    keep_cols = ["问题编号范围", "原始编号明细", "命中次数", "原始级别", "处理级别", "类别", "规则", "期望值", "实际值", "文本片段", "修复建议", "高级定位信息"]
    df = df[[c for c in keep_cols if c in df.columns]]

    show_internal_loc = st.checkbox("显示高级定位信息", value=False)
    if not show_internal_loc and "高级定位信息" in df.columns:
        df = df.drop(columns=["高级定位信息"])

    st.dataframe(df, use_container_width=True, hide_index=True)


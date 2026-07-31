from __future__ import annotations

from dataclasses import dataclass, asdict
from io import BytesIO
import re
from typing import List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING


@dataclass
class Violation:
    level: str
    category: str
    rule: str
    location: str
    expected: str
    actual: str
    snippet: str
    suggestion: str

    def to_dict(self):
        return asdict(self)


DEFAULT_RULES = {
    # body / global text rule
    "font_name": "宋体",
    "font_name_alt": "SimSun",
    "font_size_pt": "四号",
    "line_spacing_pt": 30.0,
    "first_line_indent_chars": 2.0,
    "paragraph_align": "left",
    "space_before_pt": 0.0,
    "space_after_pt": 0.0,
    "allow_bold": False,
    "allow_italic": False,
    "allow_non_black_text": False,
    "allow_underline": False,
    "allow_empty_para": False,
    "allow_leading_space": False,
    "allow_any_space": False,
    # page
    "paper": "A4",
    "margin_top_cm": 2.5,
    "margin_bottom_cm": 2.0,
    "margin_left_cm": 2.0,
    "margin_right_cm": 2.0,
    "max_pages": 100,
    # table / object
    "allow_table_row_break": True,
    "table_alignment": "center",
    "require_table_border": True,
    "table_text_align": "left",
    "table_vertical_align": "center",
    "table_first_line_indent_chars": 0.0,
    "allow_floating_shapes": False,
    "picture_alignment": "center",
    # header/footer
    "allow_header": False,
    "allow_page_number": False,
    "allow_toc": False,
    "allow_cover": False,
    # text
    "forbidden_chars": ',;:?!"\'[]{}*@#$&<>_`~\\',
    "supplier_keywords": "有限公司,有限责任公司,公司,集团,电话,邮箱,地址,联系人,法定代表人,统一社会信用代码,供应商",
    # 身份信息检测
    "pii_level": "high",  # low / medium / high
    "pii_high_sensitive_default": True,
}

PAPER_SIZES_CM = {
    "A4": (21.0, 29.7),
    "A3": (29.7, 42.0),
    "A5": (14.8, 21.0),
    "B5": (17.6, 25.0),
    "Letter": (21.59, 27.94),
}

CN_FONT_SIZES_PT = {
    "初号": 42.0,
    "小初": 36.0,
    "一号": 26.0,
    "小一": 24.0,
    "二号": 22.0,
    "小二": 18.0,
    "三号": 16.0,
    "小三": 15.0,
    "四号": 14.0,
    "小四": 12.0,
    "五号": 10.5,
    "小五": 9.0,
}


def _twips_to_cm(twips: int) -> float:
    return float(twips) * 2.54 / 1440.0


def _cm_to_twips(cm: float) -> int:
    return int(round(float(cm) / 2.54 * 1440.0))


def _resolve_font_size_pt(v) -> float:
    if isinstance(v, str):
        s = v.strip()
        if s in CN_FONT_SIZES_PT:
            return CN_FONT_SIZES_PT[s]
        try:
            return float(s)
        except Exception:
            return 14.0
    return float(v)


def _font_size_to_pt(run) -> Optional[float]:
    if run.font.size is None:
        return None
    return float(run.font.size.pt)


def _first_not_none(*vals):
    for v in vals:
        if v is not None:
            return v
    return None


def _extract_rfonts_candidates_from_xml(obj) -> List[str]:
    vals: List[str] = []
    try:
        rpr = getattr(obj, "_element", None)
        if rpr is not None:
            xml = rpr.xml
        else:
            xml = getattr(obj, "_r", None).xml if getattr(obj, "_r", None) is not None else ""
        for key in ["w:eastAsia=", "w:ascii=", "w:hAnsi=", "w:cs="]:
            parts = xml.split(key)
            if len(parts) <= 1:
                continue
            for seg in parts[1:]:
                quote = '"' if seg.startswith('"') else ("'" if seg.startswith("'") else None)
                if not quote:
                    continue
                token = seg[1:].split(quote)[0].strip()
                if token:
                    vals.append(token)
    except Exception:
        return []
    return vals


def _style_font_candidates(style) -> List[str]:
    out: List[str] = []
    if not style:
        return out
    try:
        if style.font.name:
            out.append(style.font.name)
    except Exception:
        pass
    out.extend(_extract_rfonts_candidates_from_xml(style))
    return [x for x in out if x]


def _effective_font_candidates(run, paragraph, doc: Document) -> List[str]:
    candidates: List[str] = []
    try:
        if run.font.name:
            candidates.append(run.font.name)
    except Exception:
        pass
    candidates.extend(_extract_rfonts_candidates_from_xml(run))
    candidates.extend(_style_font_candidates(getattr(run, "style", None)))
    candidates.extend(_style_font_candidates(getattr(paragraph, "style", None)))
    try:
        candidates.extend(_style_font_candidates(doc.styles["Normal"]))
    except Exception:
        pass

    seen = set()
    deduped = []
    for c in candidates:
        if c not in seen:
            seen.add(c)
            deduped.append(c)
    return deduped


def _effective_font_name(run, paragraph, doc: Document) -> Optional[str]:
    candidates = _effective_font_candidates(run, paragraph, doc)
    return candidates[0] if candidates else None


def _effective_font_size_pt(run, paragraph, doc: Document) -> Optional[float]:
    normal = None
    try:
        n = doc.styles["Normal"].font.size
        normal = float(n.pt) if n is not None else None
    except Exception:
        normal = None

    vals = [
        _font_size_to_pt(run),
        float(run.style.font.size.pt) if getattr(run, "style", None) and run.style.font.size is not None else None,
        float(paragraph.style.font.size.pt) if paragraph.style and paragraph.style.font.size is not None else None,
        normal,
    ]
    return _first_not_none(*vals)


def _effective_bool(run, paragraph, doc: Document, attr: str):
    normal = None
    try:
        normal = getattr(doc.styles["Normal"].font, attr)
    except Exception:
        normal = None
    run_style_val = getattr(getattr(run, "style", None), "font", None)
    run_style_val = getattr(run_style_val, attr) if run_style_val is not None else None
    para_style_val = getattr(paragraph.style.font, attr) if paragraph.style else None
    return _first_not_none(getattr(run.font, attr), run_style_val, para_style_val, normal, False)


def _line_spacing_to_pt(paragraph) -> Optional[float]:
    spacing = paragraph.paragraph_format.line_spacing
    if spacing is None:
        return None
    if hasattr(spacing, "pt"):
        return float(spacing.pt)
    if isinstance(spacing, (int, float)):
        return float(spacing)
    return None


def _first_line_indent_chars(paragraph, font_size_pt: float) -> Optional[float]:
    indent = paragraph.paragraph_format.first_line_indent
    if indent is None:
        return None
    return float(indent.pt) / float(font_size_pt)


def _has_forbidden_chars(text: str, chars: str) -> List[str]:
    return sorted({c for c in chars if c and c in text})


def _is_black_or_auto(run) -> bool:
    c = run.font.color
    if c is None or c.rgb is None:
        return True
    return str(c.rgb).upper() in {"000000", "AUTO"}


def _table_alignment_ok(table, expected: str) -> bool:
    if table.alignment is None:
        return False
    m = {"left": 0, "center": 1, "right": 2}
    return int(table.alignment) == m.get(expected, 1)


def _paragraph_alignment_ok(p, expected: str) -> bool:
    if p.paragraph_format.alignment is None:
        return expected == "left"
    m = {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "right": WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    return p.paragraph_format.alignment == m.get(expected, WD_ALIGN_PARAGRAPH.LEFT)


def _contains_page_field(footer) -> bool:
    xml = footer._element.xml
    return "PAGE" in xml or "w:fldSimple" in xml


def _contains_floating_shapes(doc: Document) -> bool:
    xml = doc._element.xml
    return "wp:anchor" in xml


def _count_estimated_pages(doc: Document) -> int:
    # docx无法精准分页，采用渲染分页符 + 手动分页符 + 1 的估算方式
    xml = doc._element.xml
    rendered_breaks = xml.count("w:lastRenderedPageBreak")
    manual_breaks = xml.count('w:type="page"') + xml.count("w:type='page'")
    return max(1, rendered_breaks + manual_breaks + 1)


def _check_page_setup(doc: Document, rules: dict) -> List[Violation]:
    out: List[Violation] = []
    for i, sec in enumerate(doc.sections, start=1):
        loc = f"页面设置#Section{i}"

        w_cm = _twips_to_cm(int(sec.page_width.twips))
        h_cm = _twips_to_cm(int(sec.page_height.twips))
        paper = rules.get("paper", "A4")
        target = PAPER_SIZES_CM.get(paper)
        if target:
            ok = (abs(w_cm - target[0]) <= 0.05 and abs(h_cm - target[1]) <= 0.05) or (
                abs(w_cm - target[1]) <= 0.05 and abs(h_cm - target[0]) <= 0.05
            )
            if not ok:
                out.append(Violation("warn", "页面", "纸张大小", loc, f"{paper} ({target[0]:.2f}×{target[1]:.2f} cm)", f"{w_cm:.2f} × {h_cm:.2f} cm", "", f"将纸张设置为 {paper}"))

        margin_checks = [
            ("上边距", int(sec.top_margin.twips), float(rules["margin_top_cm"])),
            ("下边距", int(sec.bottom_margin.twips), float(rules["margin_bottom_cm"])),
            ("左边距", int(sec.left_margin.twips), float(rules["margin_left_cm"])),
            ("右边距", int(sec.right_margin.twips), float(rules["margin_right_cm"])),
        ]
        for name, actual_twips, expected_cm in margin_checks:
            exp_twips = _cm_to_twips(expected_cm)
            if actual_twips != exp_twips:
                out.append(
                    Violation(
                        "warn",
                        "页面",
                        name,
                        loc,
                        f"{expected_cm:.2f} cm ({exp_twips} twips)",
                        f"{_twips_to_cm(actual_twips):.3f} cm ({actual_twips} twips)",
                        "",
                        f"将{name}调整为 {expected_cm:.2f} cm",
                    )
                )

    est_pages = _count_estimated_pages(doc)
    if est_pages > int(rules.get("max_pages", 100)):
        out.append(Violation("warn", "页面", "页数上限", "文档", f"≤ {int(rules.get('max_pages',100))} 页", f"估算约 {est_pages} 页", "", "缩减内容或拆分文档（注：为估算值）"))
    return out


def _check_header_footer_and_toc_cover(doc: Document, rules: dict) -> List[Violation]:
    out: List[Violation] = []

    front_paras = doc.paragraphs[:120]
    if not rules.get("allow_toc", False):
        toc_hit = None
        xml = doc._element.xml
        has_toc_field = (" TOC " in xml) or ('w:instr="TOC' in xml) or ("w:instrText" in xml and "TOC" in xml)
        for i, p in enumerate(front_paras, start=1):
            style_name = p.style.name if p.style else ""
            t = (p.text or "").strip()
            heading_like = bool(re.match(r"^目录\s*$", t))
            line_like = bool(re.search(r"\.{2,}\s*\d+$", t))
            if heading_like or style_name.startswith("TOC") or style_name.startswith("目录") or line_like:
                toc_hit = (i, t or style_name)
                break
        if has_toc_field or toc_hit:
            loc = f"段落#{toc_hit[0]}" if toc_hit else "文档"
            actual = f"检测到目录迹象: {toc_hit[1][:30]}" if toc_hit else "检测到 TOC 字段"
            out.append(Violation("warn", "结构", "目录", loc, "不设置目录", actual, (toc_hit[1] if toc_hit else "")[:60], "删除目录并改为正文"))

    if not rules.get("allow_cover", False):
        score = 0
        hit = None
        cover_words = ["封面", "投标文件", "技术标", "项目名称", "投标人", "编制单位"]
        for i, p in enumerate(doc.paragraphs[:35], start=1):
            t = (p.text or "").strip()
            if not t:
                continue
            local = 0
            if any(w in t for w in cover_words):
                local += 1
            if len(t) <= 18:
                local += 1
            if p.paragraph_format.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                local += 1
            if local >= 2:
                score += 1
                hit = (i, t)
        if score >= 2 and hit:
            out.append(Violation("warn", "结构", "封面", f"段落#{hit[0]}", "不设置封面", f"检测到封面迹象（评分{score}）: {hit[1][:20]}", hit[1][:60], "删除封面页"))

    for i, sec in enumerate(doc.sections, start=1):
        loc = f"页眉页脚#Section{i}"
        header_text = "".join(p.text for p in sec.header.paragraphs).strip()
        if (not rules["allow_header"]) and header_text:
            out.append(Violation("warn", "页眉", "页眉内容", loc, "无页眉文字", f"发现: {header_text[:40]}", "", "删除页眉文字"))

        if (not rules["allow_page_number"]) and _contains_page_field(sec.footer):
            out.append(Violation("warn", "页脚", "页码", loc, "无页码", "检测到 PAGE 字段", "", "删除页码"))
    return out


def _check_table_text_and_format(doc: Document, rules: dict, para_idx: dict, para_page: dict) -> List[Violation]:
    out: List[Violation] = []
    resolved_font_size = _resolve_font_size_pt(rules["font_size_pt"])

    for i, tbl in enumerate(doc.tables, start=1):
        # 表格位置锚点：取首个单元格首段落
        anchor_p = None
        try:
            anchor_p = tbl.rows[0].cells[0].paragraphs[0]
        except Exception:
            anchor_p = None
        if anchor_p is not None:
            loc = f"{_loc_para(anchor_p, para_idx, para_page)} / 表格#{i}"
        else:
            loc = f"表格#{i}"

        if not _table_alignment_ok(tbl, rules["table_alignment"]):
            out.append(Violation("warn", "表格", "整体对齐", loc, rules["table_alignment"], str(tbl.alignment), "", "将表格整体设为居中"))

        if rules["require_table_border"]:
            xml = tbl._tbl.xml
            has_border = "w:tblBorders" in xml and "w:val=\"nil\"" not in xml
            if not has_border:
                out.append(Violation("warn", "表格", "边框", loc, "有边框", "未检测到有效边框", "", "为表格添加边框"))

        for r_i, row in enumerate(tbl.rows, start=1):
            if not rules.get("allow_table_row_break", True):
                if not _row_cant_split(row):
                    out.append(
                        Violation(
                            "error",
                            "表格",
                            "跨页断行",
                            f"{loc} / 第{r_i}行",
                            "不允许行跨页断开（该行需关闭‘允许跨页断行’）",
                            "未检测到 cantSplit（该行允许跨页断行）",
                            "",
                            "在 Word 中选中该行→表格属性→行→取消‘允许跨页断行’",
                        )
                    )
            for c_i, cell in enumerate(row.cells, start=1):
                col_title = ""
                try:
                    col_title = (tbl.rows[0].cells[c_i - 1].text or "").strip().replace("\n", " ")[:12]
                except Exception:
                    col_title = ""
                cell_loc = f"{loc} / 行{r_i}列{c_i}" + (f"（列名:{col_title}）" if col_title else "")
                va = cell.vertical_alignment
                if rules.get("table_vertical_align") == "center":
                    if va is None or va != WD_ALIGN_VERTICAL.CENTER:
                        out.append(Violation("warn", "表格", "单元格垂直对齐", cell_loc, "居中", str(va), "", "设置单元格垂直居中"))

                for p in cell.paragraphs:
                    txt = (p.text or "").strip()
                    expected_align = rules.get("table_text_align", "left")
                    if txt and not _paragraph_alignment_ok(p, expected_align):
                        cn_align = {"left": "左对齐", "center": "居中", "right": "右对齐", "justify": "两端对齐"}.get(expected_align, expected_align)
                        out.append(Violation("warn", "表格", "单元格水平对齐", cell_loc, cn_align, str(p.paragraph_format.alignment), txt[:60], f"将表内文字设置为{cn_align}"))

                    fi = _first_line_indent_chars(p, resolved_font_size)
                    if fi is not None and abs(fi - float(rules.get("table_first_line_indent_chars", 0.0))) > 0.2:
                        out.append(Violation("warn", "表格", "首行缩进", cell_loc, "0 字符", f"{fi:.2f} 字符", txt[:60], "取消表格内首行缩进"))

                    out.extend(_check_runs_style(doc, p, p.runs, cell_loc, txt, rules, resolved_font_size, category="表格"))

    return out


def _check_objects(doc: Document, rules: dict) -> List[Violation]:
    out: List[Violation] = []

    if (not rules["allow_floating_shapes"]) and _contains_floating_shapes(doc):
        out.append(Violation("warn", "对象", "浮动对象", "文档", "无浮动对象", "检测到 wp:anchor", "", "改为嵌入型对象"))

    for p_i, p in enumerate(doc.paragraphs, start=1):
        has_pic = any("pic:pic" in r._r.xml or "a:blip" in r._r.xml for r in p.runs)
        if not has_pic:
            continue
        if not _paragraph_alignment_ok(p, rules["picture_alignment"]):
            out.append(Violation("warn", "图片", "对齐", f"图片段落#{p_i}", rules["picture_alignment"], str(p.paragraph_format.alignment), "", "将图片段落设为居中"))

    # 启发式：检测可能“表格转图片”情况（仅提示）
    if len(doc.tables) == 0:
        pic_count = sum(1 for p in doc.paragraphs for r in p.runs if ("pic:pic" in r._r.xml or "a:blip" in r._r.xml))
        if pic_count > 0:
            out.append(Violation("info", "对象", "疑似表格图片", "文档", "表格应使用原生表格", f"检测到 {pic_count} 张图片且无原生表格", "", "若这些图片是表格，请改为可编辑表格"))

    return out


def _row_cant_split(row) -> bool:
    try:
        tr = row._tr
        tr_pr = tr.trPr
        if tr_pr is None:
            return False
        return any(getattr(x, "tag", "").endswith("cantSplit") for x in tr_pr)
    except Exception:
        return False


def _context_around(text: str, token: str, width: int = 10) -> str:
    if not text:
        return ""
    if not token:
        return text[: min(len(text), width * 2)]
    i = text.find(token)
    if i < 0:
        return text[: min(len(text), width * 2)]
    s = max(0, i - width)
    e = min(len(text), i + len(token) + width)
    return f"{text[s:i]}【{text[i:i+len(token)]}】{text[i+len(token):e]}"


def _para_page_maps(doc: Document):
    para_idx = {}
    para_page = {}
    page = 1
    for i, p in enumerate(doc.paragraphs, start=1):
        para_idx[id(p)] = i
        para_page[id(p)] = page
        xml = p._p.xml
        page += xml.count("w:lastRenderedPageBreak") + xml.count('w:type="page"') + xml.count("w:type='page'")
    return para_idx, para_page


def _loc_para(p, para_idx: dict, para_page: dict) -> str:
    idx = para_idx.get(id(p), 0)
    pg = para_page.get(id(p), 1)
    head = (p.text or "").strip().replace("\n", " ")[:20]
    return f"约第{pg}页 / 段落#{idx} / {head}"


def _pii_hits(text: str, supplier_keywords: List[str], pii_level: str, high_sensitive_default: bool) -> List[Tuple[str, str, str]]:
    """返回[(severity, reason, token)]，同段落多命中不漏。"""
    hits: List[Tuple[str, str, str]] = []

    patterns_high = [
        (r"1[3-9]\d{9}", "手机号"),
        (r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", "邮箱"),
        (r"\b\d{17}[\dXx]\b", "身份证号"),
        (r"\b\d{18}\b", "统一社会信用代码候选"),
        (r"开户行|银行账号|账号[:：]?\s*\d{8,}", "银行信息"),
    ]
    for p, reason in patterns_high:
        for m in re.finditer(p, text):
            sev = "error" if high_sensitive_default else "warn"
            hits.append((sev, reason, m.group(0)[:30]))

    context_words = ["联系人", "法定代表人", "统一社会信用代码", "供应商", "通信地址", "联系电话", "邮箱"]
    for w in context_words:
        if w in text:
            hits.append(("warn", "身份语义词", w))

    for kw in supplier_keywords:
        if not kw:
            continue
        start = 0
        while True:
            pos = text.find(kw, start)
            if pos < 0:
                break
            if not (kw in {"公司", "集团"} and len(text) <= 6):
                sev = "info"
                if pii_level == "medium":
                    sev = "warn"
                elif pii_level == "high" and kw not in {"公司", "集团"}:
                    sev = "warn"
                hits.append((sev, "关键词命中", kw))
            start = pos + len(kw)

    rank = {"error": 3, "warn": 2, "info": 1}
    merged = {}
    for sev, reason, token in hits:
        key = (reason, token)
        if key not in merged or rank[sev] > rank[merged[key][0]]:
            merged[key] = (sev, reason, token)
    return list(merged.values())


def _font_match_ok(candidates: List[str], rules: dict) -> bool:
    expected = {str(rules.get("font_name", "")).strip(), str(rules.get("font_name_alt", "")).strip()}
    expected = {x for x in expected if x}
    if not expected:
        return True
    for c in candidates:
        cc = str(c).strip()
        if cc in expected:
            return True
    return False


def _check_runs_style(doc: Document, paragraph, runs, loc: str, txt_for_snippet: str, rules: dict, resolved_font_size: float, category: str) -> List[Violation]:
    out: List[Violation] = []
    for run in runs:
        rtxt = (run.text or "").strip()
        if not rtxt:
            continue
        font_candidates = _effective_font_candidates(run, paragraph, doc)
        rn = font_candidates[0] if font_candidates else None
        rs = _effective_font_size_pt(run, paragraph, doc)

        if not _font_match_ok(font_candidates, rules):
            out.append(Violation("warn", category, "字体", loc, f"{rules['font_name']} / {rules['font_name_alt']}", str(rn), rtxt[:60], "统一字体"))
        if rs is None or abs(rs - resolved_font_size) > 0.1:
            out.append(Violation("warn", category, "字号", loc, f"{rules['font_size_pt']} ({resolved_font_size:.1f} pt)", str(rs), rtxt[:60], "统一字号"))
        if (not rules["allow_bold"]) and bool(_effective_bool(run, paragraph, doc, "bold")):
            out.append(Violation("info", "字符样式" if category == "字体" else category, "加粗", loc, "非加粗", "加粗", rtxt[:60], "取消加粗"))
        if (not rules["allow_italic"]) and bool(_effective_bool(run, paragraph, doc, "italic")):
            out.append(Violation("info", "字符样式" if category == "字体" else category, "斜体", loc, "非斜体", "斜体", rtxt[:60], "取消斜体"))
        if (not rules["allow_underline"]) and bool(_effective_bool(run, paragraph, doc, "underline")):
            out.append(Violation("info", "字符样式" if category == "字体" else category, "下划线", loc, "无下划线", "有下划线", rtxt[:60], "取消下划线"))
        if (not rules["allow_non_black_text"]) and (not _is_black_or_auto(run)):
            out.append(Violation("warn", "字符样式" if category == "字体" else category, "字体颜色", loc, "黑色/自动", "非黑色", rtxt[:60], "改为黑色"))
    return out


def _check_paragraph_common_rules(p, txt_raw: str, txt: str, loc: str, rules: dict, resolved_font_size: float) -> List[Violation]:
    out: List[Violation] = []
    if (not rules["allow_leading_space"]) and (txt_raw.startswith(" ") or txt_raw.startswith("\u3000")):
        out.append(Violation("warn", "段落", "段首空格", loc, "不使用空格缩进", "检测到段首空格", txt[:60], "改为首行缩进设置"))

    if (not rules.get("allow_any_space", False)) and ((" " in txt_raw) or ("\u3000" in txt_raw)):
        out.append(Violation("warn", "文本", "空格", loc, "全文不允许出现半角/全角空格", "检测到空格", txt_raw[:80], "删除所有空格（含全角空格）"))

    bad = _has_forbidden_chars(txt, rules["forbidden_chars"])
    if bad:
        out.append(Violation("error", "文本", "禁用字符", loc, "仅中文/全角标点", f"发现: {' '.join(bad)}", txt[:80], "替换禁用字符"))

    if not _paragraph_alignment_ok(p, rules["paragraph_align"]):
        out.append(Violation("warn", "段落", "对齐", loc, rules["paragraph_align"], str(p.paragraph_format.alignment), txt[:60], "设置为左对齐"))

    ls_rule = p.paragraph_format.line_spacing_rule
    ls = _line_spacing_to_pt(p)
    if ls_rule is not None and ls_rule != WD_LINE_SPACING.EXACTLY:
        out.append(Violation("warn", "段落", "行距规则", loc, "固定值", str(ls_rule), txt[:60], "将行距规则设为固定值"))
    if ls is not None and abs(ls - float(rules["line_spacing_pt"])) > 0.5:
        out.append(Violation("warn", "段落", "固定行距", loc, f"{rules['line_spacing_pt']} pt", f"{ls:.1f} pt", txt[:60], "设置固定行距"))

    fi = _first_line_indent_chars(p, resolved_font_size)
    if fi is not None and abs(fi - float(rules["first_line_indent_chars"])) > 0.35:
        out.append(Violation("warn", "段落", "首行缩进", loc, f"{rules['first_line_indent_chars']} 字符", f"{fi:.2f} 字符(估算)", txt[:60], "设置首行缩进"))

    sb = p.paragraph_format.space_before.pt if p.paragraph_format.space_before else 0.0
    sa = p.paragraph_format.space_after.pt if p.paragraph_format.space_after else 0.0
    if abs(sb - float(rules["space_before_pt"])) > 0.1:
        out.append(Violation("warn", "段落", "段前距", loc, f"{rules['space_before_pt']} pt", f"{sb:.1f} pt", txt[:60], "设置段前距"))
    if abs(sa - float(rules["space_after_pt"])) > 0.1:
        out.append(Violation("warn", "段落", "段后距", loc, f"{rules['space_after_pt']} pt", f"{sa:.1f} pt", txt[:60], "设置段后距"))
    return out


def _check_body_paragraphs(doc: Document, rules: dict, para_idx: dict, para_page: dict) -> List[Violation]:
    out: List[Violation] = []
    resolved_font_size = _resolve_font_size_pt(rules["font_size_pt"])

    supplier_keywords = [x.strip() for x in str(rules.get("supplier_keywords", "")).split(",") if x.strip()]
    pii_level = str(rules.get("pii_level", "high")).lower()
    if pii_level not in {"low", "medium", "high"}:
        pii_level = "high"
    high_sensitive_default = bool(rules.get("pii_high_sensitive_default", True))

    for idx, p in enumerate(doc.paragraphs, start=1):
        txt_raw = p.text or ""
        txt = txt_raw.strip()
        loc = _loc_para(p, para_idx, para_page)

        # 分页符/分节符检查（严格）
        p_xml = p._p.xml
        # 仅将“手动分页符”视为违规；lastRenderedPageBreak 是渲染产物，不能等同人工分页
        if 'w:type="page"' in p_xml or "w:type='page'" in p_xml:
            out.append(Violation("error", "结构", "分页符", loc, "不设置分页符", "检测到手动分页符", txt[:60], "删除分页符，改为连续排版"))
        if "w:sectPr" in p_xml:
            out.append(Violation("error", "结构", "分节符", loc, "不设置分节符", "检测到分节符", txt[:60], "删除分节符，改为单一节"))

        # skip table paragraph
        in_table = False
        try:
            in_table = p._p.xpath("ancestor::w:tbl") != []
        except Exception:
            pass
        if in_table:
            continue

        if not txt:
            if not rules["allow_empty_para"]:
                out.append(Violation("warn", "段落", "空段落", loc, "不允许空段落", "空段落", "", "删除多余空段落"))
            continue

        out.extend(_check_paragraph_common_rules(p, txt_raw, txt, loc, rules, resolved_font_size))

        pii_hits = _pii_hits(txt, supplier_keywords, pii_level, high_sensitive_default)
        for sev, reason, token in pii_hits:
            out.append(
                Violation(
                    sev,
                    "暗标合规",
                    "疑似身份信息",
                    loc,
                    "不出现可识别主体信息（联系方式/证件号/供应商身份）",
                    f"命中{reason}: {token}",
                    _context_around(txt, token, 10),
                    "建议替换为匿名表述（如‘某单位’），并删除联系方式/证件号等可追溯信息",
                )
            )

        out.extend(_check_runs_style(doc, p, p.runs, loc, txt, rules, resolved_font_size, category="字体"))

    return out


def check_docx(file_bytes: bytes, rules: dict | None = None) -> List[Violation]:
    rules = {**DEFAULT_RULES, **(rules or {})}
    doc = Document(BytesIO(file_bytes))
    para_idx, para_page = _para_page_maps(doc)

    out: List[Violation] = []
    out.extend(_check_page_setup(doc, rules))
    out.extend(_check_header_footer_and_toc_cover(doc, rules))
    out.extend(_check_table_text_and_format(doc, rules, para_idx, para_page))
    out.extend(_check_objects(doc, rules))
    out.extend(_check_body_paragraphs(doc, rules, para_idx, para_page))

    return out
